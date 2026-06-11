"""
Generate BootHop Operations & Session Log document
Output: BootHopPipeline\\BootHop_OPS_Document.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

NAVY   = RGBColor(0x0D, 0x1B, 0x3E)
BLUE   = RGBColor(0x1A, 0x6F, 0xBF)
GREEN  = RGBColor(0x16, 0x75, 0x3C)
GOLD   = RGBColor(0xB8, 0x86, 0x0B)
RED    = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    # Bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A6FBF')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text).font.size = Pt(10)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BootHop')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Operations Manual & Session Log')
run.font.size = Pt(18)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generated: {datetime.now().strftime("%d %B %Y  %H:%M")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONFIDENTIAL — INTERNAL USE ONLY')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RED

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SCHEDULED TASKS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('SECTION 1 — SCHEDULED TASKS REGISTER')
body('All tasks run under the Windows SYSTEM account (ServiceAccount logon, highest privileges). '
     'The PC must be on Sleep (not Shutdown) for wake timers to fire.')

doc.add_paragraph()

TASKS = [
    {
        'name':    'BootHop-MusicGen',
        'time':    'Daily 05:00',
        'script':  'music/generate_daily_music.py',
        'purpose': 'Generates fresh background music tracks for the day\'s video content.',
        'smoke':   'python "C:\\Users\\babso\\Desktop\\BootHopPipeline\\music\\generate_daily_music.py"',
        'pass':    'New .mp3 files appear in music/output/',
        'status':  'Active',
    },
    {
        'name':    'BootHop-TikTok',
        'time':    'Daily 06:00',
        'script':  'pipeline.py',
        'purpose': 'Main TikTok POV content pipeline — generates and uploads daily short-form video.',
        'smoke':   'python "C:\\Users\\babso\\Desktop\\BootHopPipeline\\pipeline.py"',
        'pass':    'Video rendered and uploaded; .txt sidecar with YouTube/TikTok URL created.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-Cleanup',
        'time':    'Monday 06:30',
        'script':  'scripts/cleanup-monday.ps1',
        'purpose': 'Weekly cleanup — removes old temp files, logs older than 30 days, and clears render cache.',
        'smoke':   'powershell -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\cleanup-monday.ps1"',
        'pass':    'Old temp files removed; cleanup log entry written.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-Morning',
        'time':    'Daily 07:00',
        'script':  'scripts/boothop-history.ps1',
        'purpose': '"This Day in Supply Chain History" — generates and uploads a morning history video.',
        'smoke':   'powershell -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\boothop-history.ps1"',
        'pass':    'Video .mp4 created in output/; YouTube URL written to _youtube.txt.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-LinkedIn',
        'time':    'Daily 07:00',
        'script':  'scripts/linkedin-daily.ps1',
        'purpose': 'Posts a daily update to the BootHop LinkedIn company page.',
        'smoke':   'powershell -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\linkedin-daily.ps1"',
        'pass':    'Post appears on LinkedIn; log entry confirms API response 201.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-Afternoon',
        'time':    'Daily 12:30',
        'script':  'scripts/boothop-premium.ps1',
        'purpose': 'Afternoon premium reel — mid-day content push across social platforms.',
        'smoke':   'powershell -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\boothop-premium.ps1"',
        'pass':    'Reel rendered and posted; platform confirmation in log.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-Evening',
        'time':    'Daily 18:00',
        'script':  'scripts/boothop-premium.ps1',
        'purpose': 'Evening premium reel — peak engagement time content push.',
        'smoke':   'powershell -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\boothop-premium.ps1"',
        'pass':    'Reel rendered and posted; platform confirmation in log.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-GmailBlog',
        'time':    'Daily 09:00 → 13:00 (every 30 min)',
        'script':  'blog/gmail_to_blog.py',
        'purpose': 'Checks Gmail for emails with subject "New BotHop Blog Draft". Extracts body and posts to Blogger. Deduplication prevents double-posting.',
        'smoke':   'python "C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\gmail_to_blog.py"',
        'pass':    'Email found → new post appears on Blogger; processed_emails.json updated.',
        'status':  'Active',
    },
    {
        'name':    'BootHop-BlogDrop',
        'time':    'Daily 09:30',
        'script':  'blog/pdf_drop_to_blog.py',
        'purpose': 'Picks up any PDF files dropped into blog/drop/, converts to HTML, posts to Blogger, moves to blog/posted/pdfs/.',
        'smoke':   'python "C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\pdf_drop_to_blog.py"',
        'pass':    'PDF posted to Blogger; file moved to posted/pdfs/; log entry written.',
        'status':  'Active',
    },
]

# Tasks table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
for i, h in enumerate(['Task Name', 'Schedule', 'Script', 'Purpose']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr[i], '0D1B3E')

for t in TASKS:
    row = table.add_row().cells
    row[0].text = t['name']
    row[0].paragraphs[0].runs[0].bold = True
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = t['time']
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    row[2].text = t['script']
    row[2].paragraphs[0].runs[0].font.name = 'Courier New'
    row[2].paragraphs[0].runs[0].font.size = Pt(8)
    row[3].text = t['purpose']
    row[3].paragraphs[0].runs[0].font.size = Pt(9)

# Column widths
from docx.shared import Cm
widths = [Cm(4), Cm(3.5), Cm(4.5), Cm(6.5)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SMOKE TESTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 2 — SMOKE TESTS')
body('Run each command manually in PowerShell or CMD to verify the task works end-to-end. '
     'Check the PASS CRITERIA column. All scripts live in: '
     'C:\\Users\\babso\\Desktop\\BootHopPipeline\\')

doc.add_paragraph()

for t in TASKS:
    heading2(t['name'])
    p = doc.add_paragraph()
    run = p.add_run('Schedule:  ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(t['time']).font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('Purpose:   ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(t['purpose']).font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('Run command:')
    run.bold = True
    run.font.size = Pt(10)
    code(t['smoke'])

    p = doc.add_paragraph()
    run = p.add_run('Pass if:   ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(t['pass'])
    run2.font.size = Pt(10)
    run2.font.color.rgb = GREEN

    p = doc.add_paragraph()
    run = p.add_run('Log file:  ')
    run.bold = True
    run.font.size = Pt(10)
    p.add_run('C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\blog_log.txt  (blog tasks)').font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PIPELINE FOLDER STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 3 — PIPELINE FOLDER STRUCTURE')

FOLDERS = [
    ('BootHopPipeline\\',             'Root pipeline folder'),
    ('  blog\\',                      'All blog automation scripts'),
    ('    drop\\',                    'DROP PDF FILES HERE for auto-posting'),
    ('    pending\\',                 'HTML posts waiting to be sent to Blogger'),
    ('    posted\\',                  'Successfully posted HTML files + JSON metadata'),
    ('    posted\\pdfs\\',            'PDFs moved here after posting'),
    ('    archive\\',                 'Posts older than 7 days archived here'),
    ('    gmail_to_blog.py',          'Gmail "New BotHop Blog Draft" → Blogger'),
    ('    post_to_blogger.py',        'Post HTML files from pending/ to Blogger'),
    ('    pdf_drop_to_blog.py',       'PDF drop folder watcher → Blogger'),
    ('    config.json',               'Gmail + Blogger credentials (keep private)'),
    ('  scripts\\',                   'Video and content generation scripts'),
    ('    make_walkthrough_v4.py',    'Final walkthrough video (v8 = current)'),
    ('    upload_to_youtube.py',      'Upload .mp4 to YouTube via Data API v3'),
    ('    make_premium_deck.py',      'ReportLab investor deck PDF generator'),
    ('    register-tasks.ps1',        'Register all 8 pipeline tasks (run as admin)'),
    ('    register-blog-drop-task.ps1', 'Register BootHop-BlogDrop task (run as admin)'),
    ('  print\\',                     'Business card generation scripts'),
    ('    generate_cards_png.py',     'Generate 300dpi PNG/JPEG business cards'),
    ('    New_Print\\',               'Output folder for card files'),
    ('  music\\',                     'Daily music generation'),
    ('  assets\\',                    'Logos, images, brand assets'),
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(['Path', 'Description']):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr2[i], '0D1B3E')

for path, desc in FOLDERS:
    row = table2.add_row().cells
    row[0].text = path
    row[0].paragraphs[0].runs[0].font.name = 'Courier New'
    row[0].paragraphs[0].runs[0].font.size = Pt(8)
    if 'DROP' in desc or 'credentials' in desc.lower():
        row[0].paragraphs[0].runs[0].font.color.rgb = GOLD
    row[1].text = desc
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    if 'DROP' in desc:
        row[1].paragraphs[0].runs[0].bold = True
        row[1].paragraphs[0].runs[0].font.color.rgb = GOLD

table2.columns[0].width = Cm(9)
table2.columns[1].width = Cm(9.5)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WEBSITE (boothop.com) CHANGES LOG
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 4 — WEBSITE CHANGES (boothop.com) — SESSION LOG')
body('Repository: https://github.com/Daddyoba12/boothop  |  Hosted: Vercel (auto-deploy on push)')

heading2('4.1  Homepage (src/app/page.tsx)')
HOMEPAGE = [
    'Hero subtext updated: "Compliance-first logistics for same-day and cross-border delivery."',
    'Added second line: "Powered by verified travellers, couriers, and logistics partners across the UK & Europe."',
    'Added third CTA button: "Earn as a Traveller" (emerald green → /register?type=booter)',
    'Flow steps updated: "Meet" renamed to "Handoff" (Post → Match → Handoff → Deliver)',
    'Trust strip updated: replaced "95% satisfaction rate" with "GDPR & customs compliant" and "Cross-border ready"',
    'Trust items: Identity verified, Secure escrow, GDPR compliant, Customs screened, Cross-border ready, Free to join',
    'CRITICAL SEO FIX: Moved useSearchParams() out of main Suspense boundary into RegisteredRedirect child component — Googlebot now reads full page HTML instead of loading spinner (was: 19 words indexed)',
    'Added Business section: "Built for teams that move fast" with 5 compliance bullets and 4 vertical industry cards',
    'Added Use Cases grid: 6 cards (Aerospace, Legal, Medical, Diaspora, Retail, Airport Hand-Carry)',
    'Added Africa expansion hint: "Expanding: UK → Europe → Africa" pill in final CTA',
]
for item in HOMEPAGE:
    bullet(item)

heading2('4.2  How It Works (src/app/how-it-works/page.tsx)')
HIWS = [
    'Replaced YouTube iframe embed with clickable card linking to /watch page',
    'Card shows play button icon with text "After watching, you\'ll return to BootHop"',
    'Payment step corrected: "Payment released 4 hours after confirmed delivery — only if no dispute raised"',
]
for item in HIWS:
    bullet(item)

heading2('4.3  New Pages Created')
NEW_PAGES = [
    '/watch  — YouTube IFrame API page, VIDEO_ID=WrBap-JjCH4, auto-redirects to boothop.com when video ends',
    '/blog  — Fetches from Blogger RSS + shows 3 static SEO posts at top',
    '/blog/customs-clearance-services  — Full article targeting "customs clearance services" (difficulty: 0)',
    '/blog/small-business-cross-border-shipping  — Targets "cross-border shipping small business"',
    '/blog/on-board-courier-time-critical-logistics  — Targets "on-board courier / time-critical logistics"',
    '/carrier-agreement  — Legal page with downloadable carrier-agreement.pdf',
]
for item in NEW_PAGES:
    bullet(item)

heading2('4.4  Contact & Phone Fixes')
PHONES = [
    'Office number corrected everywhere: +44 115 661 2825 (was +44 115 661 282)',
    'WhatsApp number updated: 07930354325 (international: 447930354325)',
    'Vercel WHATSAPP_NUMBER env var updated via CLI',
    'src/app/api/whatsapp/route.ts fallback updated to 447930354325',
    'src/app/business/contact/page.tsx href and display text corrected',
]
for item in PHONES:
    bullet(item)

heading2('4.5  SEO Improvements')
SEO = [
    'Layout.tsx keywords expanded by 12 terms: customs clearance, on-board courier, time-critical logistics, AOG, pharmaceutical courier, verified courier, cross-border shipping UK, etc.',
    'Sitemap updated: added /blog, 3 blog posts, /business, /about, /watch, /carrier-agreement',
    'Blog listing page: static SEO posts always shown first; no empty state if Blogger is down',
    'Fixed boothop.co.uk → boothop.com link in blog post CTA',
    'how-it-works/layout.tsx: title, description, keywords, canonical, OG already in place',
]
for item in SEO:
    bullet(item)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — BLOG POSTS PUBLISHED
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 5 — BLOG POSTS PUBLISHED THIS SESSION')

POSTS = [
    {
        'title':   'Beyond the Border: Why AI is the Secret to Seamless Customs Clearance',
        'target':  'customs clearance services (Difficulty: 0, Volume: 210/mo, CPC: £5)',
        'channel': 'Blogger + /blog/customs-clearance-services on boothop.com',
        'date':    '19 May 2026',
        'labels':  'Customs & Compliance, Cross-Border Delivery, AI Logistics, BootHop',
    },
    {
        'title':   '5 Reasons Why Verified Travellers Are the Secret to Same-Day Global Shipping',
        'target':  'urgent international courier, cross-border shipping, international delivery services',
        'channel': 'Blogger (from PDF in blogfolder)',
        'date':    '19 May 2026',
        'labels':  'Verified Travellers, Same-Day Delivery, International Courier, BootHop',
    },
    {
        'title':   'On-Board Courier vs. Express Air Freight: Which Is Better For Your Urgent International Courier Needs?',
        'target':  'on-board courier, urgent international courier, OBC vs air freight',
        'channel': 'Blogger (from PDF in blogfolder)',
        'date':    '19 May 2026',
        'labels':  'On-Board Courier, Express Air Freight, Time-Critical Logistics, BootHop',
    },
    {
        'title':   'Scale Fast: The Small Business Guide to Cross-Border Shipping in 2026',
        'target':  'cross-border shipping, small business shipping, shipping for small business',
        'channel': '/blog/small-business-cross-border-shipping on boothop.com',
        'date':    '19 May 2026',
        'labels':  'Small Business, Cross-Border Shipping, B2B Logistics',
    },
    {
        'title':   'Zero to Destination: How On-Board Couriers Are Solving Time-Critical Logistics',
        'target':  'time-critical logistics, on-board courier UK, OBC delivery',
        'channel': '/blog/on-board-courier-time-critical-logistics on boothop.com',
        'date':    '19 May 2026',
        'labels':  'Time-Critical Logistics, On-Board Courier, Enterprise',
    },
]

for i, post in enumerate(POSTS, 1):
    heading2(f'{i}. {post["title"]}')
    p = doc.add_paragraph()
    p.add_run('Target keyword: ').bold = True
    p.add_run(post['target']).font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('Published to: ').bold = True
    p.add_run(post['channel']).font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run('Labels: ').bold = True
    p.add_run(post['labels']).font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — QUICK REFERENCE COMMANDS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 6 — QUICK REFERENCE COMMANDS')

heading2('Re-register all pipeline tasks (run as Administrator)')
code('powershell -ExecutionPolicy Bypass -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts\\register-tasks.ps1"')

heading2('Re-register blog drop task (run as Administrator)')
code('powershell -ExecutionPolicy Bypass -File "C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\register-blog-drop-task.ps1"')

heading2('Post PDFs from drop folder manually')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog')
code('python pdf_drop_to_blog.py')

heading2('Post HTML files from pending folder manually')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog')
code('python post_to_blogger.py')

heading2('Generate business card PNGs (300dpi)')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\print')
code('python generate_cards_png.py')

heading2('Generate investor deck PDF')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts')
code('python make_premium_deck.py')

heading2('Generate walkthrough video (v4 script → v8+ output)')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts')
code('python make_walkthrough_v4.py')

heading2('Upload video to YouTube')
code('cd C:\\Users\\babso\\Desktop\\BootHopPipeline\\scripts')
code('python upload_to_youtube.py path\\to\\video.mp4')

heading2('Deploy website to Vercel (production)')
code('cd C:\\Users\\babso\\Desktop\\boothop\\boothop')
code('vercel deploy --prod')

heading2('View all BootHop tasks in Task Scheduler')
code('Get-ScheduledTask | Where-Object { $_.TaskName -like "BootHop*" } | Format-Table TaskName, State')

heading2('Check blog posting log')
code('type C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\blog_log.txt')

heading2('How to post a new blog from PDF going forward')
body('1. Save your PDF with a descriptive filename, e.g. my-new-post.pdf')
body('2. Drop it into:  C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\drop\\')
body('3. BootHop-BlogDrop task runs at 09:30 and posts it automatically.')
body('4. PDF moves to posted\\pdfs\\ after posting.')
body('OR run manually:  python pdf_drop_to_blog.py  from the blog\\ folder.')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ENVIRONMENT & CREDENTIALS REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('SECTION 7 — ENVIRONMENT REFERENCE')
body('DO NOT share this document externally. All credentials are stored in Vercel env vars and config.json.')

CREDS = [
    ('Website URL',          'https://www.boothop.com'),
    ('GitHub repo',          'https://github.com/Daddyoba12/boothop'),
    ('Vercel project',       'tunde-olufeks-projects/boothop'),
    ('Blogger email',        'daddyoba12.obaloluwa@blogger.com'),
    ('Gmail account',        'daddyoba12@gmail.com'),
    ('Blog ID (Blogger)',    '8031835400295900689'),
    ('YouTube video ID',     'WrBap-JjCH4  (How BootHop Works walkthrough)'),
    ('Office phone',         '+44 115 661 2825'),
    ('Mobile phone',         '+44 7506 553 755'),
    ('WhatsApp number',      '07930354325  (international: 447930354325)'),
    ('Support email',        'info@boothop.com'),
    ('Business email',       'titi.olufeko@boothop.com'),
    ('Pipeline root',        'C:\\Users\\babso\\Desktop\\BootHopPipeline\\'),
    ('Website root',         'C:\\Users\\babso\\Desktop\\boothop\\boothop\\'),
    ('Config file',          'C:\\Users\\babso\\Desktop\\BootHopPipeline\\blog\\config.json'),
]

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(['Item', 'Value']):
    hdr3[i].text = h
    hdr3[i].paragraphs[0].runs[0].bold = True
    hdr3[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    hdr3[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr3[i], '0D1B3E')

for item, val in CREDS:
    row = table3.add_row().cells
    row[0].text = item
    row[0].paragraphs[0].runs[0].bold = True
    row[0].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].text = val
    row[1].paragraphs[0].runs[0].font.size = Pt(9)
    row[1].paragraphs[0].runs[0].font.name = 'Courier New'

table3.columns[0].width = Cm(5)
table3.columns[1].width = Cm(13)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\babso\Desktop\BootHopPipeline\BootHop_OPS_Document.docx'
doc.save(out)
print(f'Saved: {out}')
