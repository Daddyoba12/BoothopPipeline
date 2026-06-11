"""
BootHop Investor Deck Generator
Produces: BootHop_Investor_Deck_2026.docx + .pdf
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, subprocess, sys
from pathlib import Path

OUT_DIR = Path(r"C:\Users\babso\Desktop\BootHopPipeline\output")
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "BootHop_Investor_Deck_2026.docx"

# ── Brand colours ──────────────────────────────────────────────────────────
GREEN   = RGBColor(0x0B, 0x6E, 0x4F)
BLUE    = RGBColor(0x0B, 0x4E, 0xA6)
GOLD    = RGBColor(0xFF, 0xCC, 0x33)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DARK    = RGBColor(0x07, 0x14, 0x28)
GREY    = RGBColor(0x55, 0x55, 0x55)
LGREY   = RGBColor(0xF4, 0xF6, 0xFA)

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin    = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin   = Cm(2.2)
    s.right_margin  = Cm(2.2)

# ── Normal style ────────────────────────────────────────────────────────────
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── Helper: shade a table cell ──────────────────────────────────────────────
def shade(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    tcPr.append(shd)

# ── Helper: horizontal rule ──────────────────────────────────────────────────
def hr(doc, color='2A5298'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

# ── Helper: add coloured heading ─────────────────────────────────────────────
def heading(doc, text, size=18, bold=True, color=BLUE, before=14, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.bold           = bold
    r.font.size      = Pt(size)
    r.font.color.rgb = color
    r.font.name      = 'Calibri'
    return p

# ── Helper: bullet point ─────────────────────────────────────────────────────
def bullet(doc, text, color=None, size=11, indent=True):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(size)
        r.font.name = 'Calibri'
        if color:
            r.font.color.rgb = color
    return p

# ── Helper: body paragraph ───────────────────────────────────────────────────
def body(doc, text, size=11, color=None, bold=False, italic=False, before=2, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold      = bold
    r.italic    = italic
    if color:
        r.font.color.rgb = color
    return p

# ── Helper: slide header banner ──────────────────────────────────────────────
def slide_header(doc, number, title, subtitle=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    shade(cell, '0B4EA6')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # Slide number
    rn = p.add_run(f'  {number}  ')
    rn.bold = True; rn.font.size = Pt(11); rn.font.name = 'Calibri'
    rn.font.color.rgb = RGBColor(0xFF,0xCC,0x33)
    # Title
    rt = p.add_run(title.upper())
    rt.bold = True; rt.font.size = Pt(16); rt.font.name = 'Calibri'
    rt.font.color.rgb = WHITE
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        rs = p2.add_run(f'  {subtitle}')
        rs.font.size = Pt(11); rs.font.name = 'Calibri'
        rs.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── Helper: speaker notes box ────────────────────────────────────────────────
def notes_box(doc, text):
    hr(doc, 'CCCCCC')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Cm(0.3)
    label = p.add_run('Speaker note:  ')
    label.bold = True; label.font.size = Pt(9.5); label.font.name = 'Calibri'
    label.font.color.rgb = GREY
    note = p.add_run(text)
    note.italic = True; note.font.size = Pt(9.5); note.font.name = 'Calibri'
    note.font.color.rgb = GREY

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
cell = tbl.rows[0].cells[0]
shade(cell, '0B4EA6')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('BootHop')
r.bold = True; r.font.size = Pt(52); r.font.name = 'Calibri'
r.font.color.rgb = WHITE

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
r2 = p2.add_run('Compliance-First Distributed Logistics Infrastructure')
r2.bold = True; r2.font.size = Pt(18); r2.font.name = 'Calibri'
r2.font.color.rgb = GOLD

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(12)
r3 = p3.add_run('Same-day & cross-border logistics powered by verified movement partners,\nproduction AI customs screening, and Stripe escrow')
r3.font.size = Pt(12); r3.font.name = 'Calibri'
r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(24)
p4.paragraph_format.space_after  = Pt(40)
r4 = p4.add_run('Oluwatoyin (Titi) Olufeko  ·  Founder & CEO\ntiti.olufeko@boothop.com  ·  boothop.com  ·  2026')
r4.font.size = Pt(11); r4.font.name = 'Calibri'
r4.font.color.rgb = WHITE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# ELEVATOR PITCH
# ═══════════════════════════════════════════════════════════════════
heading(doc, 'Elevator Pitch', 14, color=BLUE, before=6, after=6)
tbl2 = doc.add_table(rows=1, cols=1)
cell2 = tbl2.rows[0].cells[0]
shade(cell2, 'F0F4FF')
ep = cell2.paragraphs[0]
ep.paragraph_format.space_before = Pt(8)
ep.paragraph_format.space_after  = Pt(8)
rr = ep.add_run(
    'BootHop turns existing passenger and commercial movement into a compliance-first distributed '
    'logistics layer for same-day, cross-border shipments. We combine verified movement partners, '
    'a production AI customs engine, manifest & inspection workflows, airport coordination, and '
    'Stripe escrow to deliver auditable, enterprise-grade urgent logistics across Africa ↔ Europe.'
)
rr.italic = True; rr.font.size = Pt(11.5); rr.font.name = 'Calibri'
rr.font.color.rgb = DARK
doc.add_paragraph().paragraph_format.space_after = Pt(4)
hr(doc)

# ═══════════════════════════════════════════════════════════════════
# SLIDES
# ═══════════════════════════════════════════════════════════════════
slides = [
    {
        "n": "01", "title": "Cover",
        "subtitle": "Compliance-First Distributed Logistics Infrastructure · 2026",
        "bullets": [
            "**Title:** BootHop",
            "**Subtitle:** Compliance-First Distributed Logistics Infrastructure",
            "**One-liner:** Same-day & cross-border logistics — verified movement partners, AI customs screening, Stripe escrow",
            "**Contact:** Oluwatoyin (Titi) Olufeko · titi.olufeko@boothop.com · boothop.com",
        ],
        "note": "Quick hello, name, role. One-line framing: show problem, solution, traction, ask. 10 slides, ~10–12 minutes.",
    },
    {
        "n": "02", "title": "The Problem",
        "subtitle": "Urgent cross-border movement is broken",
        "bullets": [
            "Urgent shipments remain expensive, fragmented, and slow",
            "Informal diaspora networks lack verification, auditability, and compliance",
            "Businesses lose time & revenue when parts, docs, or medical items are delayed",
            "Existing carriers don't fill a trusted, same-day, customs-aware gap at scale",
        ],
        "note": "Give 1 short example — engineering firm missing a flight-critical part; law firm needing originals same-day. Stress trust and compliance gap.",
    },
    {
        "n": "03", "title": "Our Solution",
        "subtitle": "A compliance-first logistics layer over existing movement",
        "bullets": [
            "BootHop converts existing passenger & commercial movement into an auditable logistics layer",
            "**Key features:** verified movement partners, open-inspection manifests, production AI customs screening, identity & escrow",
            "**Use cases:** same-day hand-carry, airport-to-airport, urgent engineering parts, legal documents, diaspora commerce",
            "Not a courier app — a compliance and trust layer enterprises can rely on",
        ],
        "note": "Emphasise 'not just a courier app' — it's an auditable logistics network built for enterprise compliance and urgent cross-border use.",
    },
    {
        "n": "04", "title": "Product & Technology",
        "subtitle": "Live platform — production AI engine — enterprise-ready stack",
        "bullets": [
            "**Production AI Compliance Engine:** country rules, restricted items, documentation checks, risk scoring",
            "**Matching & routing:** verified movement partners with airport-structured workflows",
            "**KYC, proof-of-contents, manifest generation, and Stripe escrow** on every movement",
            "**Core stack:** Next.js · Supabase · Google Maps · Stripe · Admin ops dashboard",
        ],
        "note": "Walk intake → open inspection → manifest → AI screening → verified assignment → tracked delivery → escrow release. Mention audit logs and dispute workflow.",
    },
    {
        "n": "05", "title": "Business Model",
        "subtitle": "Multiple revenue streams — consumer volume + enterprise margin",
        "bullets": [
            "**Transaction fees:** sender fee + platform commission (consumer and diaspora movements)",
            "**Enterprise premiums:** same-day, airport hand-carry, SLA packages",
            "**Recurring revenue:** priority business accounts and annual retainers",
            "**Upsell:** shipment insurance & protection, customs-tech licensing, embedded compliance APIs",
        ],
        "note": "Consumer volume feeds the pipeline; enterprise services and insurance yield higher margins and predictable retention. Mention pilot pricing (3-shipment pilot).",
    },
    {
        "n": "06", "title": "Market Opportunity",
        "subtitle": "Africa ↔ Europe corridor — diaspora + enterprise logistics",
        "bullets": [
            "Large informal diaspora flows + urgent enterprise logistics demand (Africa ↔ Europe corridor)",
            "**Addressable segments:** SME exporters, healthcare, aerospace MRO, legal/corporate, events & production",
            "**Early focus:** UK ↔ Nigeria / Ghana; scalable to EU and global hubs",
            "Billions in informal diaspora delivery flows each year — unstructured, unverified, uncompliant",
        ],
        "note": "Emphasise corridor product-market fit and explain why Africa↔Europe is an efficient place to prove compliance controls and build enterprise relationships.",
    },
    {
        "n": "07", "title": "Go-to-Market & Corridor Strategy",
        "subtitle": "Pilot → scale → partnerships",
        "bullets": [
            "**Phase 1:** Build Africa ↔ Europe gateway — control compliance & verify movement partners",
            "**B2B pilots:** 3-shipment pilot packages for engineering, healthcare, legal & e-commerce",
            "**Channels:** airport intake partners, diaspora community partnerships, enterprise procurement sales",
            "**Partnerships:** airports, insurers, customs advisors, verified courier networks",
        ],
        "note": "Describe pilot mechanics and benefits: short procurement cycle, measurable SLAs, and clear proof points for scaling enterprise and insurance relationships.",
    },
    {
        "n": "08", "title": "Traction & KPIs",
        "subtitle": "Live platform — AI engine in production — active enterprise interest",
        "bullets": [
            "Live operational platform with AI compliance engine in production",
            "Early shipments coordinated; business logistics portal deployed",
            "Organic community traction; active enterprise enquiries and pilot pipeline forming",
            "**KPIs to track:** pilots onboarded · on-time SLA % · dispute resolution time · gross transaction volume (GTV)",
        ],
        "note": "Present 3–5 concrete metrics where available. If not public, share qualitative traction: live deployments, enterprise conversations, and platform activity.",
    },
    {
        "n": "09", "title": "Team & Advisors",
        "subtitle": "Founder-led — operational depth — domain expertise",
        "bullets": [
            "**Oluwatoyin (Titi) Olufeko — Founder & CEO:** logistics, aviation systems, enterprise infra, compliance workflows",
            "**Key hires planned:** engineering lead · compliance lead · Lagos operations head · enterprise BD",
            "**Advisors:** customs & airport operations experts, insurance partners",
            "Platform built entirely by the founder — zero external development cost to date",
        ],
        "note": "Emphasise founder operational background as a core asset; show hiring plan, timeline, and how this capital accelerates team building.",
    },
    {
        "n": "10", "title": "Ask & Use of Funds",
        "subtitle": "£350,000 pre-seed · SAFE / Convertible Note · SEIS/EIS eligible",
        "bullets": [
            "**Raise:** up to £350,000 (pre-seed)",
            "**Compliance & Legal:** £50k–£80k — international legal, GDPR/NDPA, customs advisory, insurance frameworks",
            "**African Gateway Operations:** £40k–£70k — Lagos intake, inspection workflows, manifest handling, airport coordination",
            "**Trust & Verification Systems:** £25k–£50k — KYC expansion, fraud prevention, identity infrastructure",
            "**Enterprise Logistics Growth:** £50k–£100k — B2B partnerships, airport relationships, business development",
            "**Brand & Community:** £15k–£40k — diaspora partnerships, strategic digital growth, creator campaigns",
            "Seeking strategic investors with corridor, aviation, or logistics expertise",
        ],
        "note": "State runway, milestones this raise enables, and strategic value you want from investors — airport introductions, insurance relationships, customs connections. Close with clear next step: 30-minute follow-up to review pilot terms.",
    },
]

for slide in slides:
    slide_header(doc, slide["n"], slide["title"], slide.get("subtitle"))
    for b in slide["bullets"]:
        bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    notes_box(doc, slide["note"])
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════
heading(doc, 'Appendix — Optional Supporting Slides', 14, color=BLUE, before=6)
appendix = [
    ('Demo Flow Visual',   'Intake → Open Inspection → Manifest → AI Screening → Verified Assignment → Tracked Delivery → Escrow Release'),
    ('Pilot Case Study',   'Anonymised example showing time/cost improvements and SLA performance from an early shipment.'),
    ('Compliance & Security', 'KYC flow, audit logs, escrow mechanics, data handling, GDPR/NDPA compliance structure.'),
    ('Pricing & Pilot Terms', 'Pilot SOW template, 3-shipment pilot pricing, and enterprise SLA package options.'),
]
for title, desc in appendix:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f'{title}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r1.font.color.rgb = BLUE
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'
    r2.font.color.rgb = DARK

hr(doc)

# ═══════════════════════════════════════════════════════════════════
# 20-SECOND DEMO SCRIPT
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, '20-Second Investor Demo Script', 16, color=GREEN, before=6)
body(doc, 'Use this microstory for the QR landing page, investor outreach emails, and slide embeds.', italic=True, color=GREY)
hr(doc, '0B6E4F')

script = [
    ('0.0 – 3.0s', 'Logo reveal', '"BootHop — same-day cross-border logistics."', 'BH mark on brand gradient; subtle scale-up reveal.'),
    ('3.0 – 7.0s', 'Problem',     '"Urgent parts and documents get delayed, costing time and money."', 'Split: engineer + paused clock / plane on tarmac.'),
    ('7.0 – 13.0s','Solution',    '"BootHop turns existing movement into a compliance-first logistics layer — verified partners, AI-powered customs screening, open manifests, and escrow."', 'Flow animation: Inspect → Manifest → AI → Assign → Track → Escrow.'),
    ('13.0–17.0s', 'Benefit',     '"Faster, auditable, and trusted same-day delivery across Africa ↔ Europe."', 'Map with UK → Lagos route and delivered package icon.'),
    ('17.0–20.0s', 'CTA',         '"Scan to watch and save contact — links.boothop.com/titi-card."', 'Large QR + short URL + BH logo.'),
]

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
hdrs = ['Timing', 'Frame', 'Voiceover', 'Visual']
for i, h in enumerate(hdrs):
    cell = tbl3.rows[0].cells[i]
    shade(cell, '0B4EA6')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.name = 'Calibri'
    r.font.color.rgb = WHITE

for timing, frame, vo, visual in script:
    row = tbl3.add_row()
    data = [timing, frame, vo, visual]
    for i, d in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(d)
        r.font.size = Pt(9); r.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Production specs
heading(doc, 'Production Specs', 12, color=BLUE, before=8, after=4)
specs = [
    'Resolution: 1280×720 (720p) · Codec: H.264 MP4 · Bitrate: 1.5–2.5 Mbps · Target file: ≤5 MB',
    'Audio: 44.1 kHz mono · VO normalised to −3 dB · Burned-in captions for muted autoplay',
    'Fonts: Inter / Calibri · Colors: gradient #0B6E4F → #0B4EA6 · Accent: #FFCC33',
    'Music: calm, confident sting on open · light transitions · VO takes priority',
]
for s in specs:
    bullet(doc, s, size=10)

# ── Save ────────────────────────────────────────────────────────────────────
doc.save(str(DOCX_PATH))
print(f"Saved: {DOCX_PATH}")

# ── Convert to PDF via Word COM (Windows) ────────────────────────────────────
PDF_PATH = DOCX_PATH.with_suffix('.pdf')
try:
    import comtypes.client
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc_obj = word.Documents.Open(str(DOCX_PATH))
    doc_obj.SaveAs(str(PDF_PATH), FileFormat=17)  # 17 = PDF
    doc_obj.Close()
    word.Quit()
    print(f"PDF saved: {PDF_PATH}")
except Exception as e:
    print(f"PDF conversion note: {e}")
    print("Install python-comtypes or open the .docx in Word and Save As PDF.")
